#!/usr/bin/env python3
"""
catv5_core.py - VERSIÓN MULTI-IDIOMA
- Soporte para selección de idioma destino (Target Language).
- Carga dinámica de abreviaturas basada en idioma fuente.
- Integración lista para DeepL con códigos ISO.
"""

import os
import re
import csv
import json
import difflib
import shutil
import tempfile
import docx
from docx import Document
import spacy
import requests  # Necesario para DeepL

# --- CONFIGURACIÓN DE IDIOMAS (Compatibles con DeepL) ---
LANGUAGES = {
    "Spanish": "ES",
    "German": "DE",
    "Arabic": "AR",
    "Italian": "IT",
    "French": "FR",
    "Russian": "RU",
    "Polish": "PL",
    "Portuguese (Brazil)": "PT-BR",
    "Portuguese (Euro)": "PT-PT",
    "Tagalog": "TL",  # Nota: DeepL soporte limitado, verificar. Si no, Google.
    "Chinese": "ZH",
    "Indonesian": "ID",
    "Hindi": "HI",
    "Afrikaans": "AF",  # Nota: DeepL soporte limitado.
    "Dutch": "NL",
    "Japanese": "JA",
    "Korean": "KO",
    "Turkish": "TR",
    "Ukrainian": "UK",
    "English (US)": "EN-US",
    "English (UK)": "EN-GB"
}

CONFIG = {
    "PROGRESS_FILE": "progreso_cattool.json",
    "TM_FILE": "memoria_traduccion.json",
    "GLOSSARY_FILE": "terminos.csv",
    "AUTOSAVE_THRESHOLD": 10,
    "FUZZY_THRESHOLD": 0.75,
    "MODEL_NAME": "en_core_web_sm",  # Modelo Spacy para el SOURCE (Inglés)
    "DEEPL_API_KEY": "b39ef28f-d01b-4dc4-a9a8-ebbbec4499d1:fx"  # <--- REEMPLAZA ESTO
}

# Cargar spaCy
try:
    nlp = spacy.load(CONFIG["MODEL_NAME"])
except Exception:
    class DummyNLP:
        def __call__(self, text):
            class Sent:
                def __init__(self, t): self.text = t

            return type('Doc', (), {'sents': [Sent(text)]})()


    nlp = DummyNLP()


class Utils:
    ABBREVIATIONS = {}

    # Mapeo de libros comunes del español al inglés para URL de Bible Gateway
    BIBLE_BOOK_MAP = {
        # Antiguo Testamento
        "génesis": "Genesis", "éxodo": "Exodus", "levítico": "Leviticus", "números": "Numbers",
        "deuteronomio": "Deuteronomy",
        "josué": "Joshua", "jueces": "Judges", "rut": "Ruth",
        "1 samuel": "1 Samuel", "2 samuel": "2 Samuel", "1 reyes": "1 Kings", "2 reyes": "2 Kings",
        "1 crónicas": "1 Chronicles", "2 crónicas": "2 Chronicles", "esdras": "Ezra", "nehemías": "Nehemiah",
        "ester": "Esther", "job": "Job", "salmos": "Psalms", "proverbios": "Proverbs", "eclesiastés": "Ecclesiastes",
        "cantares": "Song of Solomon", "isaías": "Isaiah", "jeremías": "Jeremiah", "lamentaciones": "Lamentations",
        "ezequiel": "Ezekiel", "daniel": "Daniel", "oseas": "Hosea", "joel": "Joel", "amós": "Amos",
        "abdías": "Obadiah",
        "jonás": "Jonah", "miqueas": "Micah", "nahúm": "Nahum", "habacuc": "Habakkuk", "sofonías": "Zephaniah",
        "hageo": "Haggai", "zacarías": "Zechariah", "malaquías": "Malachi",

        # Nuevo Testamento
        "mateo": "Matthew", "marcos": "Mark", "lucas": "Luke", "juan": "John", "hechos": "Acts",
        "romanos": "Romans",
        "1 corintios": "1 Corinthians", "2 corintios": "2 Corinthians", "gálatas": "Galatians",
        "efesios": "Ephesians", "filipenses": "Philippians", "colosenses": "Colossians",
        "1 tesalonicenses": "1 Thessalonians", "2 tesalonicenses": "2 Thessalonians",
        "1 timoteo": "1 Timothy", "2 timoteo": "2 Timothy", "tito": "Titus", "filemón": "Philemon",
        "hebreos": "Hebrews", "santiago": "James",
        "1 pedro": "1 Peter", "2 pedro": "2 Peter",
        "1 juan": "1 John", "2 juan": "2 John", "3 juan": "3 John",
        "judas": "Jude", "apocalipsis": "Revelation",

        # Abreviaturas comunes en español (opcional)
        "rom": "Romans", "mt": "Matthew", "jn": "John", "hch": "Acts", "ap": "Revelation",
        "sal": "Psalms", "pr": "Proverbs", "rv": "Revelation"
    }

    @staticmethod
    def load_abbreviations(lang_code="EN"):
        """
        Carga el archivo abb_XX.csv correspondiente al idioma.
        """
        filename = f"abb_{lang_code}.csv"
        Utils.ABBREVIATIONS = {}

        if not os.path.exists(filename):
            # Fallback silencioso o print para debug
            print(f"Info: No se encontró diccionario de abreviaturas para {lang_code} ({filename})")
            return

        try:
            with open(filename, mode='r', encoding='utf-8-sig') as f:
                reader = csv.reader(f)
                for row in reader:
                    if len(row) >= 2:
                        Utils.ABBREVIATIONS[row[0].strip()] = row[1].strip()
            print(f"Abreviaturas cargadas ({lang_code}): {len(Utils.ABBREVIATIONS)}")
        except Exception as e:
            print(f"Error leyendo {filename}: {e}")

    @staticmethod
    def normalize(text):
        if not text: return ""
        return re.sub(r'\s+', ' ', text).strip().lower()

    @staticmethod
    def atomic_save(data, filepath):
        if not data: return False
        dir_name = os.path.dirname(os.path.abspath(filepath)) or "."
        try:
            with tempfile.NamedTemporaryFile("w", dir=dir_name, delete=False, encoding="utf-8") as tmp:
                json.dump(data, tmp, ensure_ascii=False, indent=2)
                tmp_name = tmp.name
            shutil.move(tmp_name, filepath)
            return True
        except Exception:
            if 'tmp_name' in locals() and os.path.exists(tmp_name):
                try:
                    os.remove(tmp_name)
                except:
                    pass
            return False

    @staticmethod
    def get_bible_url(text):
        """
        Detecta citas bíblicas, traduce el nombre del libro a inglés para el URL,
        y genera un enlace a Bible Gateway (RVR1960).
        """
        if not text: return None

        # Regex para capturar la(s) cita(s) bíblica(s) completas
        regex = r'(?:[\(\[]?)(\d?\s*[A-Za-zñáéíóúÁÉÍÓÚ\s\.]+\s+\d+:\d+(?:-\d+)?(?:(?:\s*,\s*|;\s*|\s+y\s+)\s*(?:\d?\s*[A-Za-zñáéíóúÁÉÍÓÚ\s\.]+\s+)?\d+:\d+(?:-\d+)?)*)(?:[\)\]]?)'

        match = re.search(regex, text.strip(), re.I)

        if match:
            citation_string = match.group(1).strip()

            # 1. Limpieza inicial
            citation_string = re.sub(r'^[\(\[]|[\)\]]$', '', citation_string).strip()

            # 2. Reemplazar separadores comunes por el separador de búsqueda (;)
            temp_string = re.sub(r'(?:,\s*|;\s*|\s+y\s+)', ';', citation_string)

            # 3. Traducir nombres de libros
            translated_query_parts = []
            citation_parts_regex = r'(\d?\s*[A-Za-zñáéíóúÁÉÍÓÚ\s\.]*)\s*(\d+:\d+(?:-\d+)?)'

            # Dividimos por el separador (;) para procesar cada cita
            parts = [p.strip() for p in temp_string.split(';') if p.strip()]

            for part in parts:
                match_part = re.search(citation_parts_regex, part.strip(), re.I)
                if match_part:
                    book_raw = match_part.group(1).strip().lower()
                    reference = match_part.group(2).strip()

                    # Limpieza adicional del nombre del libro (ej: quitar puntos)
                    book_clean = re.sub(r'\.', '', book_raw).strip()

                    # Traducción: Usamos el mapeo o mantenemos el texto si no se encuentra (fallback)
                    book_en = Utils.BIBLE_BOOK_MAP.get(book_clean, book_clean.title())

                    # Reconstruir la cita con el nombre en inglés
                    translated_query_parts.append(f"{book_en} {reference}")

            # Unir las citas traducidas con ';'
            query = ";".join(translated_query_parts)

            if not query: return None

            # 4. Codificación final para URL
            # El espacio ' ' debe ser %20 y los dos puntos ':' deben ser %3A
            query = query.replace(':', '%3A').replace(' ', '%20')

            BASE_URL = "https://www.biblegateway.com/passage/?search="
            VERSION = "&version=RVR1960"  # Usamos la versión completa como en el ejemplo

            return {
                "en": f"{BASE_URL}{query}{VERSION}",
                "type": "bible"
            }
        return None

    @staticmethod
    def get_egw_url(text):
        """Genera enlace a EGW Writings (Inglés) o Google."""
        if not text: return {"en": None, "type": "none"}

        # Aseguramos que las abreviaturas estén cargadas (por defecto EN si no se ha llamado antes)
        if not Utils.ABBREVIATIONS:
            Utils.load_abbreviations("EN")

        BASE_URL_EN = "https://m.egwwritings.org/en/search?query="
        GOOGLE_URL = "https://www.google.com/search?q="

        query_en = None
        is_google_fallback = False

        # 1. Cita ya abreviada { DA 12.3 }
        match_abbr = re.search(r'\{\s*([A-Za-z0-9]+)\s+(\d+(?:\.\d+)?)\s*\}', text)
        if match_abbr:
            book = match_abbr.group(1)
            ref = match_abbr.group(2)
            query_en = f"{book}+{ref}"

        # 2. Título Completo
        else:
            regex = r'^(?:[\(\[])?(.+?),\s*(?:p\.|pp\.|page|vol\.\s*\d+)?\s*(\d+(?:\.\d+)?)(?:[\)\]])?$'
            match_full = re.search(regex, text.strip())

            if match_full:
                raw_title = match_full.group(1).strip()
                ref = match_full.group(2)

                if raw_title in Utils.ABBREVIATIONS:
                    abbr = Utils.ABBREVIATIONS[raw_title]
                    query_en = f"{abbr}+{ref}"
                else:
                    is_google_fallback = True
                    clean_title = raw_title.replace(" ", "+")
                    query_en = f"{clean_title}+{ref}+Ellen+White"

        if is_google_fallback and query_en:
            return {"en": f"{GOOGLE_URL}{query_en}", "type": "google"}
        elif query_en:
            return {"en": f"{BASE_URL_EN}{query_en}", "type": "egw"}

        return {"en": None, "type": "none"}

    @staticmethod
    def get_mt_translation(text, target_lang="ES"):
        """Traduce usando DeepL API (o placeholder)."""
        if not text.strip(): return ""

        # --- CONFIGURACIÓN DEEPL ---
        API_KEY = CONFIG.get("DEEPL_API_KEY", "")
        URL = "https://api-free.deepl.com/v2/translate"

        if not API_KEY or API_KEY == "TU_CLAVE_AQUI":
            return "Error: Configura API Key en catv5_core.py"

        headers = {"Authorization": f"DeepL-Auth-Key {API_KEY}"}
        data = {
            'text': [text],
            'source_lang': 'EN',
            'target_lang': target_lang.upper()
        }

        try:
            response = requests.post(URL, headers=headers, data=data, timeout=10)
            response.raise_for_status()
            result = response.json()
            if 'translations' in result:
                return result['translations'][0]['text']
        except Exception as e:
            return f"Error MT: {e}"
        return "Error desconocido MT"


class TranslationMemory:
    def __init__(self, filepath):
        self.filepath = filepath
        self.data = {}
        self.dirty = False
        self._load_from_path(self.filepath)
        self.tm_file = filepath # Usa la ruta pasada por ProjectManager
        self._load_from_path(self.filepath)

    def _load_from_path(self, filepath):
        """Función interna que maneja la lógica de carga y resetea la TM."""
        self.data = {}  # Reset TM
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                self.data = json.load(f)

            self.tm_file = filepath
            # 🟢 MENSAJE SIN COLOR
            print(f"✅ Memoria de Traducción cargada desde: {filepath} ({len(self.data)} segmentos)")
            return True
        except FileNotFoundError:
            # ⚠️ MENSAJE SIN COLOR
            print(f"⚠️ Archivo TM no encontrado en: {filepath}. Usando memoria vacía.")
            return False
        except json.JSONDecodeError:
            # ❌ MENSAJE SIN COLOR
            print(f"❌ Error: El archivo '{filepath}' no es un JSON válido.")
            return False
        except Exception as e:
            # ❌ MENSAJE SIN COLOR
            print(f"❌ Error al cargar TM '{filepath}': {e}")
            return False

    def load_new_tm(self, filepath):
        """Método público para cargar un nuevo archivo TM desde la GUI."""
        success = self._load_from_path(filepath)
        return success

    def _load(self):
        if not os.path.exists(self.filepath): return {}
        try:
            with open(self.filepath, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return {}

    def save(self):
        if self.dirty:
            Utils.atomic_save(self.data, self.filepath)
            self.dirty = False

    def add(self, source, target):
        if not source or not target: return
        norm_src = Utils.normalize(source)
        norm_tgt = Utils.normalize(target)
        if self.data.get(norm_src) != norm_tgt:
            self.data[norm_src] = norm_tgt
            self.dirty = True

    def lookup_exact(self, source):
        return self.data.get(Utils.normalize(source))

    def lookup_fuzzy(self, source, threshold=0.75):
        norm_source = Utils.normalize(source)
        if not norm_source: return None, 0.0
        best_match = None
        best_score = 0.0
        matcher = difflib.SequenceMatcher(None, norm_source, "")

        for tm_src in self.data.keys():
            if abs(len(tm_src) - len(norm_source)) / len(norm_source) > 0.4: continue
            matcher.set_seq2(tm_src)
            if matcher.quick_ratio() < threshold: continue
            score = matcher.ratio()
            if score > best_score and score >= threshold:
                best_score = score
                best_match = self.data[tm_src]
                if best_score > 0.99: break
        return best_match, int(best_score * 100)

    def search_concordance(self, query):
        norm_query = Utils.normalize(query)
        results = []
        if not norm_query: return results
        for tm_src, tm_tgt in self.data.items():
            if norm_query in tm_src: results.append((tm_src, tm_tgt))
        return results


class Glossary:
    def __init__(self, glossary_file_path):
        self.data = {}
        self.glossary_file = glossary_file_path # Usa la ruta pasada por ProjectManager
        self._load_from_path(self.glossary_file)

    def _load_from_path(self, filepath):
        """Función interna que maneja la lógica de carga y resetea el glosario."""
        self.data = {}  # Reinicia el glosario
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    # Debe tener 2 columnas no vacías
                    if len(row) >= 2 and row[0].strip() and row[1].strip():
                        src = Utils.normalize(row[0])
                        tgt = Utils.normalize(row[1])
                        self.data[src] = tgt

            self.glossary_file = filepath
            # 🟢 MENSAJE SIN COLOR
            print(f"✅ Glosario cargado desde: {filepath} ({len(self.data)} términos)")
            return True

        except FileNotFoundError:
            # ⚠️ MENSAJE SIN COLOR
            print(f"⚠️ Archivo de glosario no encontrado en: {filepath}. Usando glosario vacío.")
            return False
        except Exception as e:
            # ❌ MENSAJE SIN COLOR
            print(f"❌ Error al cargar glosario '{filepath}': {e}")
            return False

    def check_qa(self, source, translation):
        missing = []
        src_norm = Utils.normalize(source)
        trans_norm = Utils.normalize(translation)
        for g_src, g_tgt in self.data.items():
            if re.search(r'\b' + re.escape(g_src) + r'\b', src_norm):
                if not re.search(r'\b' + re.escape(g_tgt) + r'\b', trans_norm):
                    missing.append((g_src, g_tgt))
        return missing


class ProjectManager:
    def __init__(self):
        self.docx_path = None
        self.structure = []
        self.p_idx = 0
        self.s_idx = 0
        self.tm = TranslationMemory(CONFIG["TM_FILE"])
        self.glossary = Glossary(CONFIG["GLOSSARY_FILE"])
        self.save_counter = 0

        # Idiomas
        self.source_lang = "EN"  # Fijo por ahora (EN > X)
        self.target_lang = "ES"  # Por defecto

    def set_languages(self, target_code, source_code="EN"):
        """Configura los idiomas y carga recursos correspondientes."""
        self.target_lang = target_code
        self.source_lang = source_code

        # Cargar abreviaturas del idioma FUENTE (ya que buscamos citas en el original)
        # Si el original es inglés, cargamos abb_EN.csv
        Utils.load_abbreviations(self.source_lang)

    def load_project(self, path, resume=True):
        self.docx_path = path.strip('\'"')
        loaded = False
        if resume and os.path.exists(CONFIG["PROGRESS_FILE"]):
            try:
                with open(CONFIG["PROGRESS_FILE"], 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if data.get("source_file") == os.path.abspath(self.docx_path):
                        self.structure = data["structure"]
                        self.p_idx = data.get("p_idx", 0)
                        self.s_idx = data.get("s_idx", 0)
                        # Intentar recuperar el target lang si se guardó (opcional)
                        loaded = True
            except:
                pass

        if not loaded:
            try:
                self._parse_docx()
                if os.path.exists(CONFIG["PROGRESS_FILE"]):
                    try:
                        os.remove(CONFIG["PROGRESS_FILE"])
                    except:
                        pass
                loaded = True
            except Exception as e:
                print(f"Error: {e}")
                loaded = False
        return loaded

    def _parse_docx(self):
        doc = Document(self.docx_path)
        self.structure = []
        nlp("")
        for p in doc.paragraphs:
            text = p.text.rstrip()
            if not text:
                self.structure.append({"paragraph": text, "sentences": []})
                continue
            doc_spacy = nlp(text)
            sents = [{"orig": s.text.strip(), "trad": ""} for s in doc_spacy.sents if s.text.strip()]
            self.structure.append({"paragraph": text, "sentences": sents})

    def get_current_state(self):
        attempts = 0
        max_attempts = len(self.structure) + 1
        while self.p_idx < len(self.structure) and attempts < max_attempts:
            para = self.structure[self.p_idx]
            if para["sentences"] and self.s_idx < len(para["sentences"]): break
            self.p_idx += 1
            self.s_idx = 0
            attempts += 1

        if self.p_idx >= len(self.structure): return None

        para = self.structure[self.p_idx]
        current_sent = para["sentences"][self.s_idx]

        tm_exact = self.tm.lookup_exact(current_sent["orig"])
        tm_fuzzy = self.tm.lookup_fuzzy(current_sent["orig"], CONFIG["FUZZY_THRESHOLD"])

        glossary_matches = []
        if self.glossary.data:
            src_norm = Utils.normalize(current_sent["orig"])
            for g_src, g_tgt in self.glossary.data.items():
                if g_src in src_norm: glossary_matches.append((g_src, g_tgt))

        bible_data = Utils.get_bible_url(current_sent["orig"])
        current_note = current_sent.get("note", "")

        return {
            "p_idx": self.p_idx,
            "s_idx": self.s_idx,
            "total_paras": len(self.structure),
            "total_sents": len(para["sentences"]),
            "paragraph_context": para["paragraph"],
            "sentence": current_sent,
            "orig_text": current_sent["orig"],
            "curr_trad": current_sent["trad"],
            "tm_exact": tm_exact,
            "tm_fuzzy": tm_fuzzy,
            "glossary_matches": glossary_matches,
            "bible_data": bible_data,
            "note": current_sent.get("note", "")
        }

    def merge_with_previous(self):
        """
        Une el segmento actual con el anterior dentro del mismo párrafo.
        """
        if self.p_idx >= len(self.structure): return False

        para = self.structure[self.p_idx]
        # No podemos unir si estamos en la primera oración (s_idx 0)
        if self.s_idx <= 0: return False

        # Índices
        curr_idx = self.s_idx
        prev_idx = self.s_idx - 1

        curr_sent = para["sentences"][curr_idx]
        prev_sent = para["sentences"][prev_idx]

        # 1. Unir contenido (Origen y Traducción)
        # Añadimos un espacio para separar, a menos que ya haya uno
        sep_orig = " " if not prev_sent["orig"].endswith(" ") else ""
        prev_sent["orig"] += sep_orig + curr_sent["orig"]

        # Solo unimos traducción si existe en el actual
        if curr_sent["trad"]:
            sep_trad = " " if prev_sent["trad"] and not prev_sent["trad"].endswith(" ") else ""
            prev_sent["trad"] += sep_trad + curr_sent["trad"]

        # 2. Unir Notas (si existen)
        if curr_sent.get("note"):
            prev_note = prev_sent.get("note", "")
            sep_note = " | " if prev_note else ""
            prev_sent["note"] = prev_note + sep_note + curr_sent["note"]

        # 3. Eliminar la oración actual (ya fusionada)
        del para["sentences"][curr_idx]

        # 4. Mover el cursor al segmento anterior (el fusionado)
        self.s_idx = prev_idx
        self.save_counter += 1
        return True

    def save_note(self, note_text):
        """Guarda la nota en el segmento actual."""
        if self.p_idx < len(self.structure):
            para = self.structure[self.p_idx]
            if para["sentences"] and self.s_idx < len(para["sentences"]):
                # .strip() elimina espacios vacíos. Si es "", la nota se borra.
                para["sentences"][self.s_idx]["note"] = note_text.strip()
                self.save_progress()

    # 2. Función para UNIR SEGMENTOS (Join)
    def merge_with_previous(self):
        """Une el segmento actual con el anterior."""
        if self.p_idx >= len(self.structure): return False

        para = self.structure[self.p_idx]
        if self.s_idx <= 0: return False  # No se puede unir el primero

        curr_idx = self.s_idx
        prev_idx = self.s_idx - 1

        curr_sent = para["sentences"][curr_idx]
        prev_sent = para["sentences"][prev_idx]

        # Unir Textos (Origen y Traducción)
        sep_orig = " " if not prev_sent["orig"].endswith(" ") else ""
        prev_sent["orig"] += sep_orig + curr_sent["orig"]

        if curr_sent["trad"]:
            sep_trad = " " if prev_sent["trad"] and not prev_sent["trad"].endswith(" ") else ""
            prev_sent["trad"] += sep_trad + curr_sent["trad"]

        # Unir Notas
        if curr_sent.get("note"):
            prev_note = prev_sent.get("note", "")
            sep_note = " | " if prev_note else ""
            prev_sent["note"] = prev_note + sep_note + curr_sent["note"]

        # Eliminar el segmento actual (ya fusionado)
        del para["sentences"][curr_idx]

        # Actualizar índice al anterior
        self.s_idx = prev_idx
        self.save_counter += 1
        return True

    # 3. Función de BÚSQUEDA
    def search_in_project(self, query, search_type="source"):
        """Busca texto en todo el proyecto."""
        results = []
        if not query: return results
        query = query.lower()

        for p, para in enumerate(self.structure):
            for s, sent in enumerate(para["sentences"]):
                text = sent.get("trad", "").lower() if search_type == "target" else sent["orig"].lower()
                if query in text:
                    preview = sent["orig"][:60] + "..."
                    results.append((p, s, preview))
        return results

    def go_to_segment_id(self, segment_id_str):
        """
        Intenta saltar a un segmento basado en ID visual (ej. "4.3").
        Recuerda que visual es 1-based, interno es 0-based.
        """
        try:
            parts = segment_id_str.split('.')
            if len(parts) != 2: return False

            # Convertir a índices internos (restar 1)
            target_p = int(parts[0]) - 1
            target_s = int(parts[1]) - 1

            return self.goto_segment(target_p, target_s)
        except ValueError:
            return False

    def update_translation(self, text):
        if self.p_idx >= len(self.structure): return
        para = self.structure[self.p_idx]
        if not para["sentences"]: return
        norm_text = Utils.normalize(text)
        para["sentences"][self.s_idx]["trad"] = text
        if norm_text: self.tm.add(para["sentences"][self.s_idx]["orig"], text)
        self.save_counter += 1
        if self.save_counter >= CONFIG["AUTOSAVE_THRESHOLD"]:
            self.save_progress()
            self.tm.save()
            self.save_counter = 0

    def goto_segment(self, p_idx, s_idx):
        if 0 <= p_idx < len(self.structure):
            if self.structure[p_idx]["sentences"] and 0 <= s_idx < len(self.structure[p_idx]["sentences"]):
                self.p_idx = p_idx
                self.s_idx = s_idx
                self.save_counter = 0
                return True
        return False

    def next_segment(self):
        if self.p_idx >= len(self.structure): return False
        para = self.structure[self.p_idx]
        if self.s_idx < len(para["sentences"]) - 1:
            self.s_idx += 1
            return True
        self.p_idx += 1
        self.s_idx = 0
        return self.get_current_state() is not None

    def prev_segment(self):
        if self.s_idx > 0:
            self.s_idx -= 1
        elif self.p_idx > 0:
            self.p_idx -= 1
            while self.p_idx >= 0 and not self.structure[self.p_idx]["sentences"]:
                self.p_idx -= 1
            if self.p_idx >= 0:
                self.s_idx = len(self.structure[self.p_idx]["sentences"]) - 1
            else:
                self.p_idx, self.s_idx = 0, 0

    def next_paragraph_jump(self):
        self.p_idx += 1
        self.s_idx = 0
        self.save_progress()

    def save_progress(self):
        state = {
            "source_file": os.path.abspath(self.docx_path),
            "p_idx": self.p_idx,
            "s_idx": self.s_idx,
            "structure": self.structure
        }
        Utils.atomic_save(state, CONFIG["PROGRESS_FILE"])

    def export_docx(self):
        """
        Exporta el contenido traducido recorriendo la estructura anidada (párrafos y frases).
        """
        # Aseguramos que 'docx' esté importado al inicio del archivo: import docx
        import docx

        if not self.docx_path or not self.structure:
            # Esta excepción será capturada por la GUI y mostrará el mensaje de error.
            raise FileNotFoundError("No hay un proyecto cargado o no hay contenido para exportar.")

        output_doc = docx.Document()

        # 1. Iterar sobre la lista de párrafos (self.structure)
        for para in self.structure:

            # 2. Verificar si hay frases dentro del párrafo
            sentences = para.get("sentences", [])

            # Si el párrafo no tiene frases (ej. solo imágenes o títulos vacíos), simplemente se salta.
            if not sentences:
                continue

                # Concatenar todas las frases traducidas de este párrafo en un solo string
            # Esto mantiene la cohesión dentro del párrafo, tal como estaba en el original.
            paragraph_text = ""
            for sent_data in sentences:
                # Usar la traducción ("trad") si existe, de lo contrario, usar el original ("orig")
                text_to_write = sent_data["trad"] if sent_data.get("trad") else sent_data["orig"]
                paragraph_text += text_to_write

            # 3. Agregar el párrafo completo al documento de salida
            output_doc.add_paragraph(paragraph_text)

        # 4. Definir la ruta de exportación
        base_name, ext = os.path.splitext(self.docx_path)
        output_path = f"{base_name}_traducido{ext}"

        # 5. Guardar el archivo
        output_doc.save(output_path)

        # 6. Devolver la ruta del archivo exportado (requerido por la GUI)
        return output_path