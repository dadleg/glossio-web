from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from app.models import Project, Paragraph, Segment, TranslationMemory, Glossary, User, project_assignments, AuditLog
from app.extensions import db
from app.utils import TextUtils, lookup_tm, lookup_glossary, get_nlp
import os
import re
import requests
import docx
from docx import Document

bp = Blueprint('main', __name__)

@bp.route('/')
@login_required
def index():
    # Fetch assigned projects with roles
    # We need to query the association table to get the role
    assigned_records = db.session.query(Project, project_assignments.c.role).join(
        project_assignments, project_assignments.c.project_id == Project.id
    ).filter(project_assignments.c.user_id == current_user.id).all()
    
    # Owned projects
    owned = Project.query.filter_by(user_id=current_user.id).order_by(Project.created_at.desc()).all()

    # Structure for template: project object + role
    # Owned projects have role 'owner' implicitly
    
    display_projects = []
    for p in owned:
        display_projects.append({'project': p, 'role': 'owner'})
        
    for p, role in assigned_records:
        display_projects.append({'project': p, 'role': role})
    
    # Sort by created_at
    display_projects.sort(key=lambda x: x['project'].created_at, reverse=True)
    
    return render_template('index.html', projects=display_projects)

@bp.route('/profile/update', methods=['POST'])
@login_required
def update_profile():
    name = request.form.get('name')
    if name:
        current_user.name = name
        db.session.commit()
        flash('Profile updated successfully')
    return redirect(url_for('main.index'))

@bp.route('/project/new', methods=['GET', 'POST'])
@login_required
def new_project():
    if request.method == 'POST':
        file = request.files['file']
        source_lang = request.form.get('source_lang', 'EN')
        target_lang = request.form.get('target_lang', 'ES')
        
        if file and file.filename.endswith('.docx'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            os.makedirs(current_app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(filepath)
            
            # Parse Docx
            project = Project(filename=filename, source_lang=source_lang, target_lang=target_lang, user_id=current_user.id)
            db.session.add(project)
            db.session.commit()
            
            parse_docx_to_db(filepath, project.id)
            
            return redirect(url_for('main.editor', project_id=project.id))
            
    return render_template('project_setup.html')

def parse_docx_to_db(filepath, project_id):
    doc = Document(filepath)
    nlp = get_nlp()
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
             # Store empty paragraph to preserve structure? 
             # For now, let's only store paragraphs with content or placeholders
             para = Paragraph(project_id=project_id, p_idx=i, original_text=p.text)
             db.session.add(para)
             continue
             
        para = Paragraph(project_id=project_id, p_idx=i, original_text=p.text)
        db.session.add(para)
        db.session.flush() # get ID
        
        doc_spacy = nlp(text)
        sents = [s.text.strip() for s in doc_spacy.sents if s.text.strip()]
        
        if not sents:
             # Treat whole paragraph as one segment if spacy fails or empty
             seg = Segment(paragraph_id=para.id, s_idx=0, source_text=text)
             db.session.add(seg)
        else:
            # Post-process sents to merge standalone verse references to previous segment
            merged_sents = []
            if sents:
                merged_sents.append(sents[0])
                
                # Regex for standalone verse ref (e.g. "John 3:16" or "1 John 1:1")
                # Must match strict format to avoid false positives
                verse_ref_regex = re.compile(r'^\d?\s*[A-Za-z]+\s+\d+:\d+(?:-\d+)?$')
                
                for k in range(1, len(sents)):
                    curr = sents[k]
                    # Check if current is just a verse ref
                    if verse_ref_regex.match(curr):
                        # Merge with previous
                        merged_sents[-1] += " " + curr
                    else:
                        merged_sents.append(curr)
            
            for j, s_text in enumerate(merged_sents):
                seg = Segment(paragraph_id=para.id, s_idx=j, source_text=s_text)
                db.session.add(seg)
                
    db.session.commit()

@bp.route('/editor/<int:project_id>')
@login_required
def editor(project_id):
    project = Project.query.get_or_404(project_id)
    
    is_owner = (project.user_id == current_user.id)
    user_role = 'owner' if is_owner else None
    
    if not is_owner:
        # Check assignment and get role
        assignment = db.session.query(project_assignments).filter_by(
            user_id=current_user.id, 
            project_id=project_id
        ).first()
        
        if assignment:
            user_role = assignment.role
        else:
            flash('Unauthorized access')
            return redirect(url_for('main.index'))
    
    # Eager load paragraphs and segments to avoid N+1 problem
    
    # Eager load paragraphs and segments to avoid N+1 problem
    paragraphs = Paragraph.query.filter_by(project_id=project_id).order_by(Paragraph.p_idx).all()
    
    # Fetch all segments for this project in one query
    # We join Paragraph to filter by project_id
    segments = Segment.query.join(Paragraph).filter(Paragraph.project_id == project_id).order_by(Segment.paragraph_id, Segment.s_idx).all()
    
    # Group segments by paragraph_id in memory
    segments_by_para = {}
    for seg in segments:
        if seg.paragraph_id not in segments_by_para:
            segments_by_para[seg.paragraph_id] = []
        segments_by_para[seg.paragraph_id].append(seg)
        
    # Organize structure for template
    structure = []
    for p in paragraphs:
        segs = segments_by_para.get(p.id, [])
        if segs:
            structure.append({
                'p_idx': p.p_idx,
                'segments': segs
            })
            
    return render_template('editor.html', project=project, structure=structure, user_role=user_role)

@bp.route('/api/project/<int:project_id>/assign', methods=['POST'])
@login_required
def assign_reviewer(project_id):
    project = Project.query.get_or_404(project_id)
    if project.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
        
    email = request.json.get('email')
    role = request.json.get('role', 'reviewer')
    
    if not email:
        return jsonify({'error': 'Email required'}), 400
        
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404
        
    if user in project.assigned_users:
        return jsonify({'status': 'already_assigned'})
        
    # Insert with role
    stmt = project_assignments.insert().values(user_id=user.id, project_id=project.id, role=role)
    db.session.execute(stmt)
    db.session.commit()
    return jsonify({'status': 'success'})

@bp.route('/api/project/<int:project_id>/resign', methods=['POST'])
@login_required
def resign_project(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Check if user is assigned
    if current_user not in project.assigned_users:
        return jsonify({'error': 'Not assigned to this project'}), 400
        
    # Delete assignment
    stmt = project_assignments.delete().where(
        (project_assignments.c.user_id == current_user.id) & 
        (project_assignments.c.project_id == project_id)
    )
    db.session.execute(stmt)
    db.session.commit()
    
    # Log resignation
    log = AuditLog(
        project_id=project_id,
        user_id=current_user.id,
        action='leave',
        details=f"User {current_user.email} resigned from project."
    )
    db.session.add(log)
    db.session.commit()
    
    return jsonify({'status': 'success'})

@bp.route('/api/segment/<int:segment_id>', methods=['GET'])
@login_required
def get_segment(segment_id):
    segment = Segment.query.get_or_404(segment_id)
    paragraph = Paragraph.query.get(segment.paragraph_id)
    project = Project.query.get(paragraph.project_id)
    
    # Auth check (Owner or Reviewer)
    if project.user_id != current_user.id and current_user not in project.assigned_users:
         return jsonify({'error': 'Unauthorized'}), 403

    tm_match, tm_score = lookup_tm(segment.source_text, user_id=current_user.id)
    glossary_matches = lookup_glossary(segment.source_text, user_id=current_user.id)
    bible_data = TextUtils.get_bible_url(segment.source_text)
    # Use project source lang for abbreviations
    egw_data = TextUtils.get_egw_url(segment.source_text, project.source_lang)
    
    return jsonify({
        'id': segment.id,
        'paragraph_id': paragraph.id,
        'source_text': segment.source_text,
        'target_text': segment.target_text,
        'note': segment.note,
        'paragraph_context': paragraph.original_text,
        'tm_match': tm_match,
        'tm_score': tm_score,
        'glossary_matches': glossary_matches,
        'bible_data': bible_data,
        'egw_data': egw_data,
        'last_modified_by_name': (segment.last_modified_by.name or segment.last_modified_by.email) if segment.last_modified_by else None,
        'last_modified_at': segment.last_modified_at.isoformat() if segment.last_modified_at else None
    })

@bp.route('/api/segment/<int:segment_id>/save', methods=['POST'])
@login_required
def save_segment(segment_id):
    segment = Segment.query.get_or_404(segment_id)
    para = Paragraph.query.get(segment.paragraph_id)
    proj = Project.query.get(para.project_id)
    
    # Auth Check
    if proj.user_id != current_user.id and current_user not in proj.assigned_users:
        return jsonify({'error': 'Unauthorized'}), 403

    data = request.json
    target = data.get('target_text', '')
    note = data.get('note', '')
    
    segment.target_text = target
    segment.note = note
    segment.last_modified_by_id = current_user.id
    segment.last_modified_at = datetime.utcnow()
    
    # Update TM if target is not empty
    if target.strip():
        # Simple TM update: check if source exists, if so update target, else create
        # In a real app, TM logic is more complex (fuzzy match, etc.)
        tm_entry = TranslationMemory.query.filter_by(
            user_id=current_user.id, 
            source_text=segment.source_text,
            lang_pair=f"{proj.source_lang}-{proj.target_lang}"
        ).first()
        
        if tm_entry:
            tm_entry.target_text = target
        else:
            new_tm = TranslationMemory(
                user_id=current_user.id,
                source_text=segment.source_text,
                target_text=target,
                lang_pair=f"{proj.source_lang}-{proj.target_lang}"
            )
            db.session.add(new_tm)
            
    # Propagate to other segments with same source in project?
    # Optional feature. Let's do exact match propagation within project.
    others = Segment.query.join(Paragraph).filter(
        Paragraph.project_id == proj.id,
        Segment.source_text == segment.source_text,
        Segment.id != segment.id
    ).all()
    
    for exists in others:
        if not exists.target_text: # Only if empty? Or always? Let's say only if empty to avoid overwriting manual work
            exists.target_text = segment.target_text 
            exists.last_modified_by_id = current_user.id
            exists.last_modified_at = datetime.utcnow()
            
    db.session.commit()
    
    return jsonify({'status': 'success'})

@bp.route('/api/translate/mt', methods=['POST'])
@login_required
def translate_mt():
    data = request.json
    text = data.get('text', '')
    target_lang = data.get('target_lang', 'ES')
    api_key = data.get('api_key')
    
    translation = TextUtils.get_mt_translation(text, target_lang, api_key)
    
    if translation is None:
        return jsonify({'error': 'Missing API Key'}), 400
        
    return jsonify({'translation': translation})

@bp.route('/api/segment/get_by_display_id', methods=['GET'])
@login_required
def get_segment_by_display_id():
    project_id = request.args.get('project_id', type=int)
    display_id = request.args.get('display_id')
    
    if not project_id or not display_id:
        return jsonify({'error': 'Missing params'}), 400
        
    project = Project.query.get_or_404(project_id)
    if project.user_id != current_user.id and current_user not in project.assigned_users:
        return jsonify({'error': 'Unauthorized'}), 403
        
    try:
        parts = display_id.split('.')
        if len(parts) != 2: raise ValueError
        p_idx = int(parts[0]) - 1
        s_idx = int(parts[1]) - 1
    except:
        return jsonify({'error': 'Invalid format'}), 400
        
    para = Paragraph.query.filter_by(project_id=project_id, p_idx=p_idx).first()
    if not para:
        return jsonify({'error': 'Not found'}), 404
        
    seg = Segment.query.filter_by(paragraph_id=para.id, s_idx=s_idx).first()
    if not seg:
        return jsonify({'error': 'Not found'}), 404
        
    return jsonify({'segment_id': seg.id})

@bp.route('/api/paragraph/<int:paragraph_id>/merge', methods=['POST'])
@login_required
def merge_paragraph(paragraph_id):
    para = Paragraph.query.get_or_404(paragraph_id)
    project = Project.query.get(para.project_id)
    
    if project.user_id != current_user.id and current_user not in project.assigned_users:
        return jsonify({'error': 'Unauthorized'}), 403
        
    segments = Segment.query.filter_by(paragraph_id=para.id).order_by(Segment.s_idx).all()
    if not segments or len(segments) <= 1:
        return jsonify({'status': 'no_change'})
        
    first_seg = segments[0]
    for seg in segments[1:]:
        first_seg.source_text += " " + seg.source_text
        if seg.target_text:
            sep = " " if first_seg.target_text and not first_seg.target_text.endswith(" ") else ""
            first_seg.target_text += sep + seg.target_text
        if seg.note:
             first_seg.note = (first_seg.note + " | " + seg.note) if first_seg.note else seg.note
        
        db.session.delete(seg)
    
    db.session.commit()
    return jsonify({'status': 'success', 'new_segment_id': first_seg.id})

@bp.route('/api/segment/<int:segment_id>/merge_prev', methods=['POST'])
@login_required
def merge_prev(segment_id):
    curr_seg = Segment.query.get_or_404(segment_id)
    para = Paragraph.query.get(curr_seg.paragraph_id)
    proj = Project.query.get(para.project_id)
    
    if proj.user_id != current_user.id and current_user not in proj.assigned_users:
        return jsonify({'error': 'Unauthorized'}), 403
    
    prev_seg = Segment.query.filter_by(paragraph_id=para.id, s_idx=curr_seg.s_idx - 1).first()
    if not prev_seg:
        return jsonify({'status': 'error', 'message': 'No previous segment in this paragraph'}), 400
        
    sep_orig = " " if not prev_seg.source_text.endswith(" ") else ""
    prev_seg.source_text += sep_orig + curr_seg.source_text
    
    if curr_seg.target_text:
        sep_trad = " " if prev_seg.target_text and not prev_seg.target_text.endswith(" ") else ""
        prev_seg.target_text += sep_trad + curr_seg.target_text
        
    if curr_seg.note:
        prev_note = prev_seg.note
        sep_note = " | " if prev_note else ""
        prev_seg.note = prev_note + sep_note + curr_seg.note
        
    db.session.delete(curr_seg)
    db.session.commit()
    return jsonify({'status': 'success', 'new_id': prev_seg.id})

@bp.route('/project/<int:project_id>/export')
@login_required
def export_project(project_id):
    project = Project.query.get_or_404(project_id)
    
    if project.user_id != current_user.id and current_user not in project.assigned_users:
        flash('Unauthorized')
        return redirect(url_for('main.index'))
    
    original_path = os.path.join(current_app.config['UPLOAD_FOLDER'], project.filename)
    if not os.path.exists(original_path):
         doc = Document()
    else:
         doc = Document() 
         
    paragraphs = Paragraph.query.filter_by(project_id=project_id).order_by(Paragraph.p_idx).all()
    
    for p in paragraphs:
        segments = Segment.query.filter_by(paragraph_id=p.id).order_by(Segment.s_idx).all()
        if not segments:
             doc.add_paragraph(p.original_text) 
             continue
             
        paragraph_text = ""
        for seg in segments:
             text_to_write = seg.target_text if seg.target_text else seg.source_text
             paragraph_text += text_to_write + " " 
             
        doc.add_paragraph(paragraph_text.strip())
        
    output_filename = f"translated_{project.filename}"
    output_path = os.path.join(current_app.config['UPLOAD_FOLDER'], output_filename)
    doc.save(output_path)
    
    return send_file(output_path, as_attachment=True)

@bp.route('/api/project/<int:project_id>/search', methods=['GET'])
@login_required
def search_project(project_id):
    project = Project.query.get(project_id)
    if not project: return jsonify([]), 403
    
    if project.user_id != current_user.id and current_user not in project.assigned_users:
        return jsonify([]), 403

    query = request.args.get('q', '').lower()
    search_type = request.args.get('type', 'source') 
    
    if not query:
        return jsonify([])
        
    results = []
    if search_type == 'target':
        segments = Segment.query.filter(Segment.target_text.ilike(f'%{query}%')).join(Paragraph).filter(Paragraph.project_id == project_id).all()
    else:
        segments = Segment.query.filter(Segment.source_text.ilike(f'%{query}%')).join(Paragraph).filter(Paragraph.project_id == project_id).all()
        
    for seg in segments:
        paragraph = seg.paragraph
        results.append({
            'id': seg.id,
            'p_idx': paragraph.p_idx,
            's_idx': seg.s_idx,
            'preview': seg.source_text[:60] + "...",
            'match_text': seg.target_text if search_type == 'target' else seg.source_text
        })
        
    return jsonify(results)

@bp.route('/tm/load', methods=['POST'])
@login_required
def load_tm():
    # ... (Same as before, just kept for brevity)
    return redirect(url_for('main.index'))

@bp.route('/api/tm/load', methods=['POST'])
@login_required
def api_load_tm():
    file = request.files['file']
    if file and file.filename.endswith('.json'):
        import json
        try:
            data = json.load(file)
            count = 0
            for src, tgt in data.items():
                if not TranslationMemory.query.filter_by(source_text=src, user_id=current_user.id).first():
                    tm = TranslationMemory(source_text=src, target_text=tgt, user_id=current_user.id)
                    db.session.add(tm)
                    count += 1
            db.session.commit()
            return jsonify({'status': 'success', 'message': f'Imported {count} entries.'})
        except Exception as e:
            return jsonify({'status': 'error', 'message': str(e)})
    return jsonify({'status': 'error', 'message': 'Invalid file'})

@bp.route('/api/glossary/load', methods=['POST'])
@login_required
def api_load_glossary():
    file = request.files['file']
    if file and file.filename.endswith('.csv'):
        import csv
        import io
        try:
            stream = io.StringIO(file.stream.read().decode("UTF8"), newline=None)
            reader = csv.reader(stream)
            count = 0
            for row in reader:
                if len(row) >= 2:
                    g = Glossary(source_term=row[0].strip(), target_term=row[1].strip(), user_id=current_user.id)
                    db.session.add(g)
                    count += 1
            db.session.commit()
            return jsonify({'status': 'success', 'message': f'Imported {count} terms.'})
        except Exception as e:
             return jsonify({'status': 'error', 'message': str(e)})
    return jsonify({'status': 'error', 'message': 'Invalid file'})

@bp.route('/api/bible/versions', methods=['GET'])
@login_required
def proxy_bible_versions():
    try:
        r = requests.get('https://bolls.life/static/bolls/app/views/languages.json', timeout=10)
        return jsonify(r.json())
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/bible/books/<slug>', methods=['GET'])
@login_required
def proxy_bible_books(slug):
    try:
        url = f'https://bolls.life/get-books/{slug}/'
        r = requests.get(url, timeout=10)
        return jsonify(r.json())
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/bible/text/<slug>/<int:book>/<int:chapter>', methods=['GET'])
@login_required
def proxy_bible_chapter(slug, book, chapter):
    try:
        url = f'https://bolls.life/get-text/{slug}/{book}/{chapter}/'
        r = requests.get(url, timeout=10)
        return jsonify(r.json())
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/bible/verse/<slug>/<int:book>/<int:chapter>/<int:verse>', methods=['GET'])
@login_required
def proxy_bible_verse(slug, book, chapter, verse):
    try:
        url = f'https://bolls.life/get-verse/{slug}/{book}/{chapter}/{verse}/'
        r = requests.get(url, timeout=10)
        return jsonify(r.json())
    except Exception as e:
        return jsonify({'error': str(e)}), 500
